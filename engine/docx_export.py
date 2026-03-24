"""
TMOS13 DOCX Export

Converts markdown deliverable body to .docx using python-docx.
Returns raw bytes for streaming response.
"""
import io
import re
import logging
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger("tmos13.docx_export")


def markdown_to_docx(title: str, markdown_body: str, metadata: dict | None = None) -> bytes:
    """
    Convert markdown to .docx bytes.

    Supports:
    - # / ## / ### headings
    - **bold** and *italic* inline formatting
    - Bullet lists (- or * prefix)
    - Paragraphs separated by blank lines
    - Horizontal rules (---)

    Args:
        title: Document title (used in header)
        markdown_body: Markdown content
        metadata: Optional dict with date, pack, visitor_name, etc.

    Returns:
        Raw bytes of the .docx file
    """
    doc = Document()
    metadata = metadata or {}

    # ── Document metadata header ──
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run(title)
    run.bold = True
    run.font.size = Pt(18)

    # Metadata line
    meta_parts = []
    if metadata.get("date"):
        meta_parts.append(metadata["date"])
    if metadata.get("pack"):
        meta_parts.append(f"Pack: {metadata['pack']}")
    if metadata.get("visitor_name"):
        meta_parts.append(f"Visitor: {metadata['visitor_name']}")
    if meta_parts:
        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(" | ".join(meta_parts))
        meta_run.italic = True
        meta_run.font.size = Pt(9)

    # Divider
    doc.add_paragraph("─" * 60)

    # ── Parse markdown body ──
    lines = markdown_body.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)
            i += 1
            continue

        # Regular paragraph — collect consecutive non-empty, non-special lines
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line or next_line.startswith("#") or next_line.startswith("- ") or \
               next_line.startswith("* ") or next_line in ("---", "***", "___") or \
               next_line.startswith("```"):
                break
            para_lines.append(next_line)
            i += 1

        para_text = " ".join(para_lines)
        p = doc.add_paragraph()
        _add_formatted_text(p, para_text)

    # ── Write to bytes ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _add_formatted_text(paragraph, text: str):
    """Add text with **bold** and *italic* inline formatting to a paragraph."""
    # Split on bold and italic markers
    # Process bold first (**text**), then italic (*text*)
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*)")

    last_end = 0
    for match in pattern.finditer(text):
        # Add text before this match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        if match.group(2):  # Bold (**text**)
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # Italic (*text*)
            run = paragraph.add_run(match.group(3))
            run.italic = True

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])
